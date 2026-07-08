"""
文件格式转换核心模块
支持 Word/Excel/PPT/PDF/图片 等格式之间的互相转换
依赖: win32com, openpyxl, python-docx, pdf2docx, PyPDF2, PIL, reportlab
"""

import os
import threading
from pathlib import Path


class FileConverter:
    """文件格式转换器
    
    提供多种文件格式之间的转换功能
    """

    # 支持的转换类型
    CONVERSIONS = {
        "word_to_pdf": {
            "name": "Word 转 PDF",
            "source_ext": [".doc", ".docx"],
            "target_ext": ".pdf",
            "description": "将Word文档转换为PDF格式"
        },
        "excel_to_pdf": {
            "name": "Excel 转 PDF",
            "source_ext": [".xls", ".xlsx"],
            "target_ext": ".pdf",
            "description": "将Excel表格转换为PDF格式"
        },
        "ppt_to_pdf": {
            "name": "PPT 转 PDF",
            "source_ext": [".ppt", ".pptx"],
            "target_ext": ".pdf",
            "description": "将PowerPoint演示文稿转换为PDF格式"
        },
        "pdf_to_word": {
            "name": "PDF 转 Word",
            "source_ext": [".pdf"],
            "target_ext": ".docx",
            "description": "将PDF文档转换为可编辑的Word文档"
        },
        "excel_to_word": {
            "name": "Excel 转 Word",
            "source_ext": [".xls", ".xlsx"],
            "target_ext": ".docx",
            "description": "将Excel数据转换为Word表格文档"
        },
        "word_to_excel": {
            "name": "Word 转 Excel",
            "source_ext": [".doc", ".docx"],
            "target_ext": ".xlsx",
            "description": "提取Word中的表格数据到Excel"
        },
        "image_to_pdf": {
            "name": "图片 转 PDF",
            "source_ext": [".jpg", ".jpeg", ".png", ".bmp", ".gif", ".tiff"],
            "target_ext": ".pdf",
            "description": "将图片（支持多张）合并为PDF文件"
        },
        "pdf_merge": {
            "name": "PDF 合并",
            "source_ext": [".pdf"],
            "target_ext": ".pdf",
            "description": "将多个PDF文件合并为一个"
        },
        "pdf_split": {
            "name": "PDF 拆分",
            "source_ext": [".pdf"],
            "target_ext": ".pdf",
            "description": "将PDF文件按页拆分为多个文件"
        },
    }

    def __init__(self):
        """初始化文件转换器"""
        pass

    def get_available_conversions(self):
        """获取所有可用的转换类型
        
        Returns:
            dict: 转换类型字典
        """
        return self.CONVERSIONS

    def convert(self, conversion_type, source_paths, target_path, progress_callback=None):
        """执行文件转换
        
        Args:
            conversion_type: 转换类型标识
            source_paths: 源文件路径列表
            target_path: 目标文件路径
            progress_callback: 进度回调函数
            
        Returns:
            tuple: (成功标志, 消息)
        """
        converter_map = {
            "word_to_pdf": self.word_to_pdf,
            "excel_to_pdf": self.excel_to_pdf,
            "ppt_to_pdf": self.ppt_to_pdf,
            "pdf_to_word": self.pdf_to_word,
            "excel_to_word": self.excel_to_word,
            "word_to_excel": self.word_to_excel,
            "image_to_pdf": self.image_to_pdf,
            "pdf_merge": self.pdf_merge,
            "pdf_split": self.pdf_split,
        }

        func = converter_map.get(conversion_type)
        if not func:
            return False, f"不支持的转换类型: {conversion_type}"

        try:
            return func(source_paths, target_path, progress_callback)
        except Exception as e:
            return False, f"转换失败: {str(e)}"

    def _report_progress(self, callback, current, total, message=""):
        """报告进度
        
        Args:
            callback: 回调函数
            current: 当前进度
            total: 总数
            message: 附加消息
        """
        if callback:
            callback(current, total, message)

    def word_to_pdf(self, source_paths, target_path, progress_callback=None):
        """Word转PDF
        
        使用Microsoft Word COM组件将Word文档转换为PDF
        """
        import win32com.client

        total = len(source_paths)
        word_app = None

        try:
            word_app = win32com.client.Dispatch("Word.Application")
            word_app.Visible = False

            for i, src in enumerate(source_paths):
                self._report_progress(progress_callback, i, total, f"正在转换: {Path(src).name}")
                doc = word_app.Documents.Open(os.path.abspath(src))
                # 17 = wdFormatPDF
                doc.SaveAs(os.path.abspath(target_path), FileFormat=17)
                doc.Close()

            self._report_progress(progress_callback, total, total, "转换完成")
            return True, f"成功将 {total} 个Word文档转换为PDF"
        finally:
            if word_app:
                word_app.Quit()

    def excel_to_pdf(self, source_paths, target_path, progress_callback=None):
        """Excel转PDF
        
        使用Microsoft Excel COM组件将Excel表格转换为PDF
        """
        import win32com.client

        total = len(source_paths)
        excel_app = None

        try:
            excel_app = win32com.client.Dispatch("Excel.Application")
            excel_app.Visible = False

            for i, src in enumerate(source_paths):
                self._report_progress(progress_callback, i, total, f"正在转换: {Path(src).name}")
                wb = excel_app.Workbooks.Open(os.path.abspath(src))
                # 0 = xlTypePDF
                wb.ExportAsFixedFormat(0, os.path.abspath(target_path))
                wb.Close()

            self._report_progress(progress_callback, total, total, "转换完成")
            return True, f"成功将 {total} 个Excel文件转换为PDF"
        finally:
            if excel_app:
                excel_app.Quit()

    def ppt_to_pdf(self, source_paths, target_path, progress_callback=None):
        """PPT转PDF
        
        使用Microsoft PowerPoint COM组件将演示文稿转换为PDF
        """
        import win32com.client

        total = len(source_paths)
        ppt_app = None

        try:
            ppt_app = win32com.client.Dispatch("PowerPoint.Application")
            ppt_app.Visible = True

            for i, src in enumerate(source_paths):
                self._report_progress(progress_callback, i, total, f"正在转换: {Path(src).name}")
                pres = ppt_app.Presentations.Open(os.path.abspath(src), WithWindow=False)
                # 32 = ppFixedFormatTypePDF
                pres.SaveAs(os.path.abspath(target_path), 32)
                pres.Close()

            self._report_progress(progress_callback, total, total, "转换完成")
            return True, f"成功将 {total} 个PPT文件转换为PDF"
        finally:
            if ppt_app:
                ppt_app.Quit()

    def pdf_to_word(self, source_paths, target_path, progress_callback=None):
        """PDF转Word
        
        使用pdf2docx库将PDF转换为可编辑的Word文档
        """
        from pdf2docx import Converter

        total = len(source_paths)

        for i, src in enumerate(source_paths):
            self._report_progress(progress_callback, i, total, f"正在转换: {Path(src).name}")
            cv = Converter(src)
            cv.convert(target_path)
            cv.close()

        self._report_progress(progress_callback, total, total, "转换完成")
        return True, f"成功将 {total} 个PDF文件转换为Word文档"

    def excel_to_word(self, source_paths, target_path, progress_callback=None):
        """Excel转Word
        
        读取Excel数据并写入Word表格
        """
        import openpyxl
        from docx import Document
        from docx.shared import Pt

        total = len(source_paths)

        for i, src in enumerate(source_paths):
            self._report_progress(progress_callback, i, total, f"正在转换: {Path(src).name}")
            wb = openpyxl.load_workbook(src)
            doc = Document()

            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                doc.add_heading(sheet_name, level=1)

                rows = ws.max_row
                cols = ws.max_column

                if rows == 0 or cols == 0:
                    doc.add_paragraph("（空表）")
                    continue

                table = doc.add_table(rows=rows, cols=cols, style='Table Grid')

                for r in range(1, rows + 1):
                    for c in range(1, cols + 1):
                        cell = ws.cell(row=r, column=c)
                        value = str(cell.value) if cell.value is not None else ""
                        table.cell(r - 1, c - 1).text = value

                doc.add_paragraph()

            doc.save(target_path)

        self._report_progress(progress_callback, total, total, "转换完成")
        return True, f"成功将 {total} 个Excel文件转换为Word文档"

    def word_to_excel(self, source_paths, target_path, progress_callback=None):
        """Word转Excel
        
        提取Word文档中的表格数据到Excel
        """
        from docx import Document
        import openpyxl

        total = len(source_paths)
        wb = openpyxl.Workbook()

        for i, src in enumerate(source_paths):
            self._report_progress(progress_callback, i, total, f"正在转换: {Path(src).name}")
            doc = Document(src)
            sheet_name = f"文件{i+1}"
            ws = wb.create_sheet(sheet_name)

            row_offset = 1
            for table in doc.tables:
                for r, row in enumerate(table.rows):
                    for c, cell in enumerate(row.cells):
                        ws.cell(row=r + row_offset, column=c + 1, value=cell.text)
                row_offset += len(table.rows) + 2

        # 删除默认Sheet
        if "Sheet" in wb.sheetnames:
            del wb["Sheet"]

        wb.save(target_path)

        self._report_progress(progress_callback, total, total, "转换完成")
        return True, f"成功从 {total} 个Word文档提取表格到Excel"

    def image_to_pdf(self, source_paths, target_path, progress_callback=None):
        """图片转PDF
        
        将一张或多张图片合并为PDF文件
        """
        from PIL import Image

        total = len(source_paths)
        images = []

        for i, src in enumerate(source_paths):
            self._report_progress(progress_callback, i, total, f"正在处理: {Path(src).name}")
            img = Image.open(src)
            if img.mode != 'RGB':
                img = img.convert('RGB')
            images.append(img)

        if not images:
            return False, "没有找到可转换的图片"

        first = images[0]
        rest = images[1:] if len(images) > 1 else []

        first.save(target_path, "PDF", save_all=True, append_images=rest)

        self._report_progress(progress_callback, total, total, "转换完成")
        return True, f"成功将 {total} 张图片合并为PDF"

    def pdf_merge(self, source_paths, target_path, progress_callback=None):
        """PDF合并
        
        将多个PDF文件合并为一个
        """
        from PyPDF2 import PdfMerger

        total = len(source_paths)
        merger = PdfMerger()

        for i, src in enumerate(source_paths):
            self._report_progress(progress_callback, i, total, f"正在合并: {Path(src).name}")
            merger.append(src)

        merger.write(target_path)
        merger.close()

        self._report_progress(progress_callback, total, total, "合并完成")
        return True, f"成功合并 {total} 个PDF文件"

    def pdf_split(self, source_paths, target_path, progress_callback=None):
        """PDF拆分
        
        将PDF文件按页拆分为多个文件
        """
        from PyPDF2 import PdfReader, PdfWriter

        total = len(source_paths)
        output_dir = os.path.dirname(target_path)
        base_name = Path(target_path).stem
        count = 0

        for i, src in enumerate(source_paths):
            self._report_progress(progress_callback, i, total, f"正在拆分: {Path(src).name}")
            reader = PdfReader(src)
            num_pages = len(reader.pages)

            for page_num in range(num_pages):
                writer = PdfWriter()
                writer.add_page(reader.pages[page_num])
                out_file = os.path.join(output_dir, f"{base_name}_第{page_num + 1}页.pdf")
                with open(out_file, "wb") as f:
                    writer.write(f)
                count += 1

        self._report_progress(progress_callback, total, total, "拆分完成")
        return True, f"成功拆分 {total} 个PDF文件，共生成 {count} 个文件"

    @staticmethod
    def check_office_available():
        """检查Microsoft Office是否可用
        
        Returns:
            dict: 各组件的可用状态
        """
        result = {"word": False, "excel": False, "ppt": False}
        try:
            import win32com.client
            for app_name, key in [("Word.Application", "word"),
                                  ("Excel.Application", "excel"),
                                  ("PowerPoint.Application", "ppt")]:
                try:
                    app = win32com.client.Dispatch(app_name)
                    result[key] = True
                    app.Quit()
                except Exception:
                    result[key] = False
        except ImportError:
            pass
        return result

    @staticmethod
    def get_default_target_path(source_path, target_ext):
        """根据源文件路径生成默认目标路径
        
        Args:
            source_path: 源文件路径
            target_ext: 目标扩展名
            
        Returns:
            str: 默认目标路径
        """
        p = Path(source_path)
        return str(p.parent / f"{p.stem}_转换{target_ext}")
